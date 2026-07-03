import os
import PyPDF2
import docx

def extract_text_from_pdf(file_path):
    """
    Extracts raw text from a PDF file using PyPDF2.
    """
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF file {file_path}: {e}")
    return text

def extract_text_from_docx(file_path):
    """
    Extracts raw text from a DOCX file using python-docx.
    """
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        print(f"Error reading DOCX file {file_path}: {e}")
    return text

def extract_text(file_path):
    """
    Detects file extension and extracts text accordingly.
    """
    _, ext = os.path.splitext(file_path.lower())
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        # Fallback to general text reading if it's plain text
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except Exception as e:
            print(f"Unsupported file format or reading error for {file_path}: {e}")
            return ""

if __name__ == "__main__":
    print("Resume parser loaded successfully.")
